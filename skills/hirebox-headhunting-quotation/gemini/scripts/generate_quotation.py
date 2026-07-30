"""Generate a role-specific Hirebox quotation while keeping the appendix fixed."""

from __future__ import annotations

import argparse
import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import BaseDocTemplate, Frame, PageBreak, PageTemplate, Paragraph, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
DOCX_TEMPLATE = ASSETS / "headhunting-quotation-template.docx"
LOGO_PATH = ASSETS / "HIREBOX_LOGO_PNG.png"
FONT = "MicrosoftYaHei"
FONT_BOLD = "MicrosoftYaHeiBold"
NAVY = "#0B2545"
BLUE = "#2E74B5"
STANDARD_SERVICE = "岗位需求沟通、候选人寻访、初步筛选、候选人推荐、面试协调、录用协助及入职跟进。\nRole briefing, candidate sourcing, preliminary screening, candidate presentation, interview coordination, offer support and onboarding follow-up."


def amount(value):
    if not isinstance(value, (int, float)) or value <= 0:
        raise ValueError("Salary values must be positive numbers in THB.")
    return float(value)


def tier_rate(annual_salary):
    if annual_salary <= 300_000:
        return 10, 60
    if annual_salary <= 900_000:
        return 15, 60
    if annual_salary <= 1_800_000:
        return 20, 90
    return 25, 90


def validate_and_prepare(raw):
    required = ["client_name", "project_name", "position_title", "headcount", "work_location", "key_requirements", "salary_basis"]
    missing = [key for key in required if not str(raw.get(key, "")).strip()]
    if missing:
        raise ValueError("Missing required fields: " + ", ".join(missing))
    basis = raw["salary_basis"]
    if basis.get("type") not in {"client_budget", "market_research"}:
        raise ValueError("salary_basis.type must be client_budget or market_research.")
    if basis.get("period") not in {"monthly", "annual"}:
        raise ValueError("salary_basis.period must be monthly or annual.")
    if basis["type"] == "market_research" and not basis.get("market_sources"):
        raise ValueError("market_research requires at least one market_sources entry.")
    low, high = amount(basis.get("min_thb")), amount(basis.get("max_thb"))
    if high < low:
        raise ValueError("salary_basis.max_thb must be greater than or equal to min_thb.")
    multiplier = 12 if basis["period"] == "monthly" else 1
    annual_low, annual_high = low * multiplier, high * multiplier
    low_rate, _ = tier_rate(annual_low)
    high_rate, _ = tier_rate(annual_high)
    annual_range = f"THB {annual_low:,.0f}" if annual_low == annual_high else f"THB {annual_low:,.0f} - THB {annual_high:,.0f}"
    rate = f"年薪的 {low_rate}% / {low_rate}% of annual salary" if low_rate == high_rate else f"年薪的 {low_rate}% - {high_rate}% / {low_rate}% - {high_rate}% of annual salary"
    source_cn = "客户薪资预算" if basis["type"] == "client_budget" else "市场岗位调研薪资范围"
    source_en = "client salary budget" if basis["type"] == "client_budget" else "market-researched salary range"
    final_note = "最终费率以候选人确认年薪对应区间为准。 / The final rate follows the candidate's confirmed annual-salary band." if low_rate != high_rate else ""
    raw["quotation_date"] = raw.get("quotation_date") or date.today().isoformat()
    raw["service_content"] = raw.get("service_content") or STANDARD_SERVICE
    raw["rate_text"] = f"基于{source_cn}，预估年薪 {annual_range}，预估适用服务费率为 {rate}。 / Based on the {source_en}, estimated annual salary is {annual_range} and the estimated applicable service fee rate is {rate}. {final_note}".strip()
    return raw


def set_cell(cell, value):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(str(value))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)


def replace_text(paragraph, old, new):
    if old in paragraph.text:
        text = paragraph.text.replace(old, new)
        paragraph.clear()
        paragraph.add_run(text)


def build_docx(data, output):
    doc = Document(DOCX_TEMPLATE)
    if len(doc.tables) < 4:
        raise RuntimeError("Quotation template structure is incomplete.")
    appendix_before = [table._tbl.xml for table in doc.tables[2:]]
    main, role = doc.tables[:2]
    for row, value in zip(main.rows, [data["client_name"], data["project_name"], data["service_content"], data["quotation_date"], "泰铢（THB） / Thai Baht (THB)"]):
        set_cell(row.cells[1], value)
    for row, value in zip(role.rows, [data["position_title"], data["headcount"], data["work_location"], data["key_requirements"], data["rate_text"]]):
        set_cell(row.cells[1], value)
    for paragraph in doc.paragraphs:
        replace_text(paragraph, "[签署日期 / Date]", data["quotation_date"])
    if appendix_before != [table._tbl.xml for table in doc.tables[2:]]:
        raise RuntimeError("Refusing to modify the fixed quotation-standard appendix.")
    doc.save(output)


def register_fonts():
    pdfmetrics.registerFont(TTFont(FONT, r"C:\Windows\Fonts\msyh.ttc"))
    pdfmetrics.registerFont(TTFont(FONT_BOLD, r"C:\Windows\Fonts\msyhbd.ttc"))


def p(text, style):
    return Paragraph(str(text).replace("\n", "<br/>"), style)


def table_style(header=False, labels=False):
    commands = [("GRID", (0, 0), (-1, -1), 0.65, colors.HexColor("#7F8C99")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 7), ("RIGHTPADDING", (0, 0), (-1, -1), 7), ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]
    if header:
        commands.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")))
    if labels:
        commands.append(("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F4F7")))
    return TableStyle(commands)


def kv(rows, body):
    table = Table([[p(k, body), p(v, body)] for k, v in rows], colWidths=[2.08 * inch, 4.42 * inch])
    table.setStyle(table_style(labels=True))
    return table


def draw_page(canvas_obj, doc):
    canvas_obj.saveState()
    canvas_obj.drawImage(str(LOGO_PATH), inch, 10.42 * inch, width=1.58 * inch, height=0.41 * inch, mask="auto")
    canvas_obj.setFillColor(colors.HexColor(NAVY))
    canvas_obj.setFont(FONT_BOLD, 9.5)
    canvas_obj.drawRightString(7.5 * inch, 10.62 * inch, "海钡人力泰国有限公司")
    canvas_obj.setFont(FONT_BOLD, 8.5)
    canvas_obj.drawRightString(7.5 * inch, 10.45 * inch, "HIREBOX CO., LTD.")
    canvas_obj.setStrokeColor(colors.HexColor("#D6DEE8"))
    canvas_obj.line(inch, 10.34 * inch, 7.5 * inch, 10.34 * inch)
    canvas_obj.setFillColor(colors.HexColor("#555555"))
    canvas_obj.setFont(FONT, 7.5)
    canvas_obj.drawCentredString(4.25 * inch, 0.34 * inch, "8th Floor, S-Metro Building, 725 Sukhumvit Road, Khlong Tan Nuea, Watthana, Bangkok 10110 | johnny@hirebox.com.cn")
    canvas_obj.restoreState()


class QuoteDoc(BaseDocTemplate):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="main")
        self.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=draw_page)])


def build_pdf(data, output):
    register_fonts()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName=FONT, fontSize=9.3, leading=12.3, textColor=colors.HexColor("#222222"), spaceAfter=4)
    title = ParagraphStyle("title", parent=body, fontName=FONT_BOLD, fontSize=22, leading=28, alignment=TA_CENTER, textColor=colors.HexColor(NAVY), spaceAfter=3)
    subtitle = ParagraphStyle("subtitle", parent=body, fontSize=11, leading=14, alignment=TA_CENTER, textColor=colors.HexColor("#555555"), spaceAfter=8)
    heading = ParagraphStyle("heading", parent=body, fontName=FONT_BOLD, fontSize=13.2, leading=17, textColor=colors.HexColor(BLUE), spaceBefore=7, spaceAfter=4)
    appendix_title = ParagraphStyle("appendix_title", parent=title, fontSize=18, leading=22, spaceAfter=2)
    appendix_subtitle = ParagraphStyle("appendix_subtitle", parent=subtitle, fontSize=9.5, leading=12, spaceAfter=4)
    appendix_heading = ParagraphStyle("appendix_heading", parent=heading, fontSize=12.5, leading=15, spaceBefore=4, spaceAfter=2)
    table_body = ParagraphStyle("table", parent=body, fontSize=8.7, leading=11.3, spaceAfter=0)
    table_head = ParagraphStyle("tablehead", parent=table_body, fontName=FONT_BOLD, alignment=TA_CENTER, textColor=colors.HexColor(NAVY))
    note = ParagraphStyle("note", parent=body, fontSize=8.1, leading=9.8, spaceAfter=0)
    sign = ParagraphStyle("sign", parent=body, fontSize=9.5, leading=14.5, alignment=TA_RIGHT, spaceAfter=2)
    rate_rows = [
        [p("候选人年薪区间\nCandidate Annual Salary Range", table_head), p("质保期\nGuarantee Period", table_head), p("服务费标准\nService Fee", table_head)],
        [p("不超过 THB 300,000\nUp to THB 300,000", table_body), p("60 天\n60 days", table_body), p("年薪的 10%\n10% of annual salary", table_body)],
        [p("THB 300,000 至 THB 900,000\nTHB 300,000 to THB 900,000", table_body), p("60 天\n60 days", table_body), p("年薪的 15%\n15% of annual salary", table_body)],
        [p("THB 900,000 至 THB 1,800,000\nTHB 900,000 to THB 1,800,000", table_body), p("90 天\n90 days", table_body), p("年薪的 20%\n20% of annual salary", table_body)],
        [p("超过 THB 1,800,000\nAbove THB 1,800,000", table_body), p("90 天\n90 days", table_body), p("年薪的 25%\n25% of annual salary", table_body)],
    ]
    rate_table = Table(rate_rows, colWidths=[2.4 * inch, 1.25 * inch, 2.85 * inch], repeatRows=1)
    rate_table.setStyle(table_style(header=True))
    story = [
        p("猎头服务报价单", title), p("Headhunting Service Quotation", subtitle),
        kv([("客户名称 / Client", data["client_name"]), ("项目名称 / Project", data["project_name"]), ("服务内容 / Service Scope", data["service_content"]), ("报价日期 / Quotation Date", data["quotation_date"]), ("报价币种 / Currency", "泰铢（THB） / Thai Baht (THB)")], table_body), Spacer(1, 6),
        p("岗位需求与报价 / Role Requirements and Quotation", heading),
        kv([("招聘岗位 / Position", data["position_title"]), ("招聘人数 / Headcount", data["headcount"]), ("工作地点 / Work Location", data["work_location"]), ("核心需求 / Key Requirements", data["key_requirements"]), ("预估使用服务费率 / Estimated Applicable Service Fee Rate", data["rate_text"])], table_body),
        p(f"<b>报价单位 / Quotation Issuer:</b> HIREBOX CO., LTD. / 海钡人力泰国有限公司", sign), p(f"<b>签署日期 / Date:</b> {data['quotation_date']}", sign), PageBreak(),
        p("报价标准附件", appendix_title), p("Appendix: Quotation Standards", appendix_subtitle), p("本附件为海钡人力固定的猎头服务报价与付款标准。 / This appendix sets out HIREBOX's standard headhunting service fee and payment terms.", body),
        p("服务费率标准 / Service Fee Schedule", appendix_heading), rate_table,
        p("<b>备注 / Note:</b> 服务费以候选人确认入职岗位的年度薪酬总额为基数，按对应年薪区间适用服务费率计算。 / The service fee is calculated on the candidate's total annual remuneration for the confirmed role, using the applicable rate for the relevant annual-salary band.", note), Spacer(1, 4),
        p("付款方式 / Payment Terms", appendix_heading),
        kv([("合同签署预付款\nContract Signing Advance", "双方合同签署后，客户向海钡人力支付 THB 20,000 预付款；该预付款可在候选人入职后应付服务费中抵扣，另有书面约定的除外。\nAfter both parties sign the contract, the client shall pay HIREBOX an advance of THB 20,000. The advance may be offset against service fees due after the candidate starts, unless otherwise agreed in writing."), ("候选人入职付款\nCandidate Start Date Payment", "候选人正式入职后，客户支付对应猎头服务费的 50%。\nUpon the candidate's official start date, the client shall pay 50% of the applicable headhunting service fee."), ("质保期结束付款\nPayment After Guarantee Period", "候选人通过适用质保期后，客户支付剩余对应猎头服务费的 50%。\nAfter the candidate completes the applicable guarantee period, the client shall pay the remaining 50% of the applicable headhunting service fee.")], table_body),
        p("<b>备注 / Note:</b> 本报价所列服务费均未含税。开具泰国增值税发票时，客户应按开票日适用税率另行支付增值税。预扣税、印花税及其他依法产生的税费，按泰国适用法律法规及双方签署的正式合同承担、申报和处理。 / All quoted service fees are exclusive of taxes. For a Thai VAT invoice, the client shall pay VAT in addition at the rate applicable on the invoice date. Withholding tax, stamp duty and other legally applicable taxes or charges shall be borne, filed and handled under Thai law and the definitive agreement signed by the parties.", note), Spacer(1, 3),
        p("服务说明 / Service Notes", appendix_heading), p("<b>质保与递补 / Guarantee and Replacement:</b> 质保期为整个岗位的质保期。质保期内如候选人离职或未通过质保期，海钡人力可根据双方确认的岗位条件免费提供一次递补候选人；替补候选人入职后，视为该岗位质保期当即结束，客户应支付剩余尾款。 / The guarantee period applies to the role as a whole. If a candidate resigns or does not complete the guarantee period, HIREBOX may provide one replacement candidate free of charge based on the agreed role requirements. Once the replacement candidate starts, the guarantee period for that role ends immediately and the client shall pay the outstanding balance.", note),
    ]
    QuoteDoc(str(output), pagesize=letter, leftMargin=inch, rightMargin=inch, topMargin=1.10 * inch, bottomMargin=0.78 * inch, title="Hirebox Headhunting Service Quotation", author="HIREBOX CO., LTD.").build(story)


def filename_part(value):
    cleaned = re.sub(r'[<>:"/\\|?*]+', "_", str(value)).strip()
    return cleaned[:40] or "quotation"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    data = validate_and_prepare(json.loads(args.input.read_text(encoding="utf-8")))
    args.output_dir.mkdir(parents=True, exist_ok=True)
    stem = f"海钡人力_猎头服务报价单_{filename_part(data['client_name'])}_{filename_part(data['position_title'])}"
    docx_path, pdf_path = args.output_dir / f"{stem}.docx", args.output_dir / f"{stem}.pdf"
    build_docx(data, docx_path)
    build_pdf(data, pdf_path)
    print(json.dumps({"docx": str(docx_path), "pdf": str(pdf_path), "estimated_rate": data["rate_text"]}, ensure_ascii=False))


if __name__ == "__main__":
    main()
