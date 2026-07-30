"""Build a Hirebox candidate recommendation report from a JSON payload."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLUE = "1F4EAB"
NAVY = "1B114C"
LIGHT_BLUE = "EAF1FF"
TEXT = "202124"
GREY = "6B7280"
GREEN = "1F7A5C"
AMBER = "B7791F"
DEFAULT_LOGO_PATH = Path(__file__).resolve().parent.parent / "assets" / "HIREBOX_LOGO_PNG.png"
CONTACT_REPLACEMENT = "[候选人联系方式已隐去]"
CONTACT_PATTERNS = [
    re.compile(r"[\w.+-]+@[\w-]+(?:\.[\w-]+)+"),
    re.compile(r"(微信|WeChat|Line|LINE|电话|手机|邮箱|Email|E-mail)\s*[:：]\s*[^\n\r;；|]+", re.IGNORECASE),
]


def sanitize_candidate_contacts(value):
    if isinstance(value, str):
        text = value
        for pattern in CONTACT_PATTERNS:
            text = pattern.sub(CONTACT_REPLACEMENT, text)
        return text
    if isinstance(value, list):
        return [sanitize_candidate_contacts(item) for item in value]
    if isinstance(value, dict):
        return {key: sanitize_candidate_contacts(item) for key, item in value.items()}
    return value


def default_pdf_path(output_path):
    return str(Path(output_path).with_suffix(".pdf"))


def resolve_logo_path(payload):
    logo_path = payload.get("logo_path")
    if logo_path and Path(logo_path).exists():
        return logo_path
    if DEFAULT_LOGO_PATH.exists():
        return str(DEFAULT_LOGO_PATH)
    return None


def export_pdf(docx_path, pdf_path):
    docx = Path(docx_path).resolve()
    pdf = Path(pdf_path).resolve()
    pdf.parent.mkdir(parents=True, exist_ok=True)

    script = f"""
$ErrorActionPreference = 'Stop'
$docx = '{str(docx).replace("'", "''")}'
$pdf = '{str(pdf).replace("'", "''")}'
$wpsApi = Get-ChildItem -Path "$env:LOCALAPPDATA\\Kingsoft\\WPS Office" -Recurse -Filter wpsapi.dll -ErrorAction SilentlyContinue |
    Sort-Object FullName -Descending |
    Select-Object -First 1
if ($wpsApi) {{
    & regsvr32 /s $wpsApi.FullName
}}
$progIds = @('Kwps.Application', 'KWPS.Application', 'Word.Application')
$lastError = $null
foreach ($progId in $progIds) {{
    try {{
        $app = New-Object -ComObject $progId
        $app.Visible = $false
        try {{
            $doc = $app.Documents.Open($docx)
            $doc.ExportAsFixedFormat($pdf, 17)
            $doc.Close($false)
        }} finally {{
            $app.Quit()
        }}
        Write-Output $progId
        exit 0
    }} catch {{
        $lastError = $_.Exception.Message
    }}
}}
throw $lastError
"""
    with tempfile.NamedTemporaryFile("w", suffix=".ps1", delete=False, encoding="utf-8-sig") as handle:
        script_path = Path(handle.name)
        handle.write(script)
    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", str(script_path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        if pdf.exists():
            return True
    except Exception:
        pass
    finally:
        script_path.unlink(missing_ok=True)

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(pdf.parent),
                str(docx),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        converted = pdf.parent / (docx.stem + ".pdf")
        if converted != pdf and converted.exists():
            converted.replace(pdf)
        return True

    raise RuntimeError("No layout-faithful PDF converter is available. Install/register WPS Writer, Microsoft Word, or LibreOffice/soffice.")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E2F3", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_borders(table, color="D9E2F3"):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, color=color)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_run(run, size=10.5, bold=False, color=TEXT):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", size=10.5, bold=False, color=TEXT, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color)
    return p


def add_section_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    style_run(r, size=15, bold=True, color=NAVY)
    if subtitle:
        add_paragraph(doc, subtitle, size=9.5, color=GREY, space_after=8)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run("• " + str(item))
        style_run(r, size=10, color=TEXT)


def add_kv_table(doc, rows, key_fill=LIGHT_BLUE):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(3.3)
        cells[1].width = Cm(12.0)
        set_cell_shading(cells[0], key_fill)
        for idx, text in enumerate((key, value)):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            style_run(r, size=9.5, bold=(idx == 0), color=NAVY if idx == 0 else TEXT)
    set_table_borders(table)


def add_matrix_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        r = cell.paragraphs[0].add_run(str(header))
        style_run(r, size=9.5, bold=True, color="FFFFFF")
    for row_values in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            color = TEXT
            if idx == 1:
                value = str(text)
                if value == "高度匹配":
                    color = GREEN
                elif value == "部分匹配":
                    color = AMBER
                elif value in {"待验证", "不匹配"}:
                    color = GREY
            r = p.add_run(str(text))
            style_run(r, size=8.8, bold=(idx < 2), color=color)
    set_table_borders(table)


def add_score_cards(doc, cards):
    if not cards:
        return
    table = doc.add_table(rows=1, cols=len(cards))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, card in enumerate(cards):
        key, value = card
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY if i == 0 else LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(key) + "\n" + str(value))
        style_run(r, size=9.5, bold=True, color="FFFFFF" if i == 0 else NAVY)
    set_table_borders(table)


def reset_header_footer(section, logo_path=None, footer_text=None):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for paragraph in section.header.paragraphs:
        paragraph._element.clear_content()
    for paragraph in section.footer.paragraphs:
        paragraph._element.clear_content()
    if logo_path and Path(logo_path).exists():
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.add_run().add_picture(str(logo_path), width=Cm(3.6))
    if footer_text:
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer_p.add_run(footer_text)
        style_run(r, size=8.5, color=GREY)


def add_content_section(doc, section):
    add_section_title(doc, section["title"], section.get("subtitle"))
    for para in section.get("paragraphs", []):
        if isinstance(para, str):
            add_paragraph(doc, para)
        else:
            add_paragraph(
                doc,
                para.get("text", ""),
                bold=para.get("bold", False),
                color=para.get("color", TEXT),
                size=para.get("size", 10.5),
            )
    if section.get("bullets"):
        add_bullets(doc, section["bullets"])
    table = section.get("table")
    if table:
        if table.get("type") == "kv":
            add_kv_table(doc, table.get("rows", []))
        elif table.get("type") == "matrix":
            add_matrix_table(doc, table.get("headers", []), table.get("rows", []))


def build(payload):
    payload = sanitize_candidate_contacts(payload)
    doc = Document()
    first_section = doc.sections[0]
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.75)
        section.right_margin = Cm(1.75)
        section.header_distance = Cm(0.7)
        section.footer_distance = Cm(0.7)
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)

    logo_path = resolve_logo_path(payload)
    footer_text = payload.get("footer", "Hirebox 海钡人力 | Confidential Candidate Recommendation Report")
    cover = payload["cover"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    add_paragraph(doc, cover.get("title", "人才推荐报告"), size=26, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_paragraph(doc, cover.get("subtitle", "Candidate Recommendation Report"), size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)
    add_paragraph(doc, cover.get("role", ""), size=15, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_paragraph(doc, cover.get("candidate", ""), size=12, color=TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)
    add_kv_table(doc, cover.get("rows", []))

    doc.add_section(WD_SECTION.NEW_PAGE)
    for section in doc.sections:
        reset_header_footer(section, logo_path=logo_path, footer_text=footer_text)

    add_score_cards(doc, payload.get("score_cards", []))
    for section_payload in payload.get("sections", []):
        add_content_section(doc, section_payload)

    # Ensure first section margins remain set after section creation.
    first_section.top_margin = Cm(1.8)
    Path(payload["output_path"]).parent.mkdir(parents=True, exist_ok=True)
    doc.save(payload["output_path"])
    pdf_path = payload.get("pdf_output_path") or default_pdf_path(payload["output_path"])
    export_pdf(payload["output_path"], pdf_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("payload", help="Path to JSON payload")
    args = parser.parse_args()
    with open(args.payload, "r", encoding="utf-8") as f:
        payload = json.load(f)
    build(payload)
    print(payload["output_path"])


if __name__ == "__main__":
    main()
