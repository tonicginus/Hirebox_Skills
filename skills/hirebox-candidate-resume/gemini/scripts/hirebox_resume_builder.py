"""Build HireBox-style candidate resume DOCX from a JSON payload.

This script handles layout only. The calling agent must extract facts, write
concise content, and run source consistency checks.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


NAVY = "1B114C"
BLUE = "3264E8"
ICE = "EEF5FF"
LAV = "F4F1FF"
TEXT = "202124"
GREY = "667085"
LINE = "D7E0F0"


def set_run(run, size=10.5, bold=False, color=TEXT):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border(cell, color=LINE, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            border(cell)


def para(container, text="", size=10.5, bold=False, color=TEXT, align=None, after=5, before=0):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    set_run(r, size=size, bold=bold, color=color)
    return p


def section_title(doc, title, subtitle=None):
    p = para(doc, title, size=15, bold=True, color=NAVY, after=3, before=12)
    p.paragraph_format.keep_with_next = True
    if subtitle:
        para(doc, subtitle, size=9, color=GREY, after=7)


def bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run("• " + str(text))
    set_run(r, size=10, color=TEXT)


def kv_table(doc, rows, widths=(3.0, 5.3, 3.0, 5.3)):
    table = doc.add_table(rows=0, cols=4)
    table.autofit = False
    for row in rows:
        values = list(row)
        if len(values) == 2:
            values = [values[0], values[1], "", ""]
        cells = table.add_row().cells
        for idx, value in enumerate(values[:4]):
            cells[idx].width = Cm(widths[idx])
            if idx % 2 == 0 and value:
                shade(cells[idx], ICE)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, size=9.2, bold=(idx % 2 == 0), color=NAVY if idx % 2 == 0 else TEXT)
    style_table(table)


def simple_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        if widths:
            cell.width = Cm(widths[i])
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(header))
        set_run(r, size=9.2, bold=True, color="FFFFFF")
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if widths:
                cells[i].width = Cm(widths[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, size=8.8, bold=(i == 0), color=TEXT)
    style_table(table)


def add_badges(doc, badges):
    if not badges:
        return
    table = doc.add_table(rows=1, cols=len(badges))
    table.autofit = False
    for i, badge in enumerate(badges):
        label = badge.get("label", "")
        value = badge.get("value", "")
        fill = badge.get("fill", [NAVY, ICE, LAV, ICE][i % 4])
        cell = table.rows[0].cells[i]
        cell.width = Cm(4.2)
        shade(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(label) + "\n" + str(value))
        set_run(r, size=9.3, bold=True, color=NAVY if fill != NAVY else "FFFFFF")
    style_table(table)


def setup_header_footer(doc, logo_path=None, footer_text="HireBox 海钡人力 | 候选人简历 | Confidential"):
    logo = Path(logo_path) if logo_path else None
    for section in doc.sections:
        section.top_margin = Cm(1.55)
        section.bottom_margin = Cm(1.35)
        section.left_margin = Cm(1.65)
        section.right_margin = Cm(1.65)
        section.header_distance = Cm(0.65)
        section.footer_distance = Cm(0.65)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p._element.clear_content()
        for p in section.footer.paragraphs:
            p._element.clear_content()
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if logo and logo.exists():
            hp.add_run().add_picture(str(logo), width=Cm(3.2))
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run(footer_text)
        set_run(r, size=8.2, color=GREY)


def build(payload):
    include_contacts = payload.get("include_contacts", True)
    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)

    first = doc.sections[0]
    first.top_margin = Cm(1.7)
    first.bottom_margin = Cm(1.4)
    first.left_margin = Cm(1.65)
    first.right_margin = Cm(1.65)

    para(doc, payload.get("title", "海钡人力候选人简历"), size=26, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=6, before=10)
    para(doc, payload.get("subtitle", "HireBox Candidate Profile"), size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    para(doc, payload.get("headline", ""), size=15, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

    add_badges(doc, payload.get("badges", []))

    overview = payload.get("overview_rows", [])
    contacts = payload.get("contact_rows", [])
    if include_contacts and contacts:
        overview = overview + contacts
    if overview:
        section_title(doc, "一、候选人概览")
        kv_table(doc, overview)

    for section in payload.get("sections", []):
        section_title(doc, section["title"], section.get("subtitle"))
        for item in section.get("paragraphs", []):
            para(doc, item)
        for item in section.get("bullets", []):
            bullet(doc, item)
        table = section.get("table")
        if table:
            simple_table(doc, table.get("headers", []), table.get("rows", []), table.get("widths"))

    doc.add_section(WD_SECTION.CONTINUOUS)
    setup_header_footer(doc, payload.get("logo_path"), payload.get("footer", "HireBox 海钡人力 | 候选人简历 | Confidential"))
    out = Path(payload["output_path"])
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("payload", help="JSON payload path")
    args = parser.parse_args()
    with open(args.payload, "r", encoding="utf-8") as f:
        payload = json.load(f)
    print(build(payload))


if __name__ == "__main__":
    main()
